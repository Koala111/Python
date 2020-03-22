from docx import Document
document = Document()

paragraph = document.add_paragraph('Lorem ipsum dolor sit amet.')
prior_paragraph = paragraph.insert_paragraph_before('Lorem ipsum')

document.add_heading('这个是二级目录', level=1)  #一级目录
document.add_heading('这个是三级目录', level=2)  #二级目录

table = document.add_table(rows=2, cols=2)
cell = table.cell(0, 1)
cell.text = 'parrot, possibly dead'
row = table.rows[1]
row.cells[0].text = 'Foo bar to you.'
row.cells[1].text = 'And a hearty foo bar to you too sir!'
for row in table.rows:
    for cell in row.cells:
        print(cell.text)

document.add_page_break() # 强制到下一页

document.save('F:\\SDV.docx')
print("word文档生成成功，请到“F:\\”目录 里面查阅")